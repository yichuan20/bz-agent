#!/usr/bin/env python3
"""
create-doc.py — convert markdown-flavoured text to a Word (.docx) document.

The agent writes content in simple markdown, then calls this script to produce
a properly-formatted DOCX file.

Usage:
  # Content from a file
  python3 bzcode/scripts/create-doc.py --content-file /tmp/draft.md --out /path/to/report.docx

  # Content passed inline (small docs)
  python3 bzcode/scripts/create-doc.py --content "# Title\n\nParagraph text." --out report.docx

Supported markdown:
  # H1  ## H2  ### H3  #### H4
  - bullet  * bullet
  | col1 | col2 |  (tables — first row becomes header)
  **bold**  *italic*  (inline, within paragraphs)
  Plain paragraphs

Output:
  Prints JSON: { "ok": true, "path": "/abs/path/to/report.docx", "paragraphs": N }
"""
import argparse
import json
import re
import sys
from pathlib import Path


def _apply_inline(run_text: str, para):
    """Parse **bold** and *italic* markers and add runs to para."""
    import docx.oxml.ns as ns
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
    for m in pattern.finditer(run_text):
        if m.group(2):          # **bold**
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):        # *italic*
            run = para.add_run(m.group(3))
            run.italic = True
        elif m.group(4):        # plain
            para.add_run(m.group(4))


def _parse_table(lines):
    """Return list-of-lists from markdown table lines."""
    rows = []
    for line in lines:
        if re.match(r'^\|[-| :]+\|$', line.strip()):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def markdown_to_docx(content: str, out_path: Path) -> int:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    # Compact default margins
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(90)
        section.right_margin  = Pt(90)

    lines    = content.splitlines()
    i        = 0
    para_cnt = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            p = doc.add_heading(text, level=level)
            para_cnt += 1
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            doc.add_paragraph('─' * 40)
            i += 1
            continue

        # ── Table block ───────────────────────────────────────────────────────
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            rows = _parse_table(table_lines)
            if rows:
                cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=cols)
                tbl.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = cell_text
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                doc.add_paragraph('')
                para_cnt += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        m = re.match(r'^[\-\*\+]\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            _apply_inline(m.group(1), p)
            para_cnt += 1
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        m = re.match(r'^\d+\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            _apply_inline(m.group(1), p)
            para_cnt += 1
            i += 1
            continue

        # ── Empty line → spacing ──────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc.add_paragraph()
        _apply_inline(stripped, p)
        para_cnt += 1
        i += 1

    doc.save(out_path)
    return para_cnt


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert markdown to Word DOCX.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--content",      help="Markdown content string.")
    group.add_argument("--content-file", help="Path to a markdown file to convert.")
    parser.add_argument("--out", required=True, help="Output .docx path.")
    args = parser.parse_args()

    if args.content_file:
        cf = Path(args.content_file)
        if not cf.exists():
            print(json.dumps({"error": f"content file not found: {cf}"}))
            sys.exit(1)
        content = cf.read_text(encoding="utf-8")
    else:
        content = args.content.replace("\\n", "\n")

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if not out_path.suffix.lower() == ".docx":
        out_path = out_path.with_suffix(".docx")

    try:
        para_cnt = markdown_to_docx(content, out_path)
    except Exception as exc:
        print(json.dumps({"error": f"could not create document: {exc}"}))
        sys.exit(1)

    print(json.dumps({"ok": True, "path": str(out_path.resolve()), "paragraphs": para_cnt}))


if __name__ == "__main__":
    main()
